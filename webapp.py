import io
import numpy as np
from PIL import Image
import streamlit as st
from streamlit_autorefresh import st_autorefresh
import requests
from fpdf import FPDF
from huggingface_hub import hf_hub_download
import tensorflow as tf
import os
import json
import time
from types import MethodType
from tensorflow.keras.applications.mobilenet_v2 import preprocess_input

# ==========================
# TRANSLATION FALLBACK (Google Translate unofficial endpoint)
# ==========================
LANG_CODE_MAP = {
    "English": "en","Hindi": "hi","Bengali": "bn","Tamil": "ta","Telugu": "te",
    "Kannada": "kn","Malayalam": "ml","Marathi": "mr","Gujarati": "gu",
    "Punjabi": "pa","Odia": "or","Urdu": "ur",
}

def translate_with_google(text, target_language_name):
    code = LANG_CODE_MAP.get(target_language_name, "en")
    if code=="en": return text
    try:
        url="https://translate.googleapis.com/translate_a/single"
        params={"client":"gtx","sl":"en","tl":code,"dt":"t","q":text}
        r=requests.get(url,params=params,timeout=15); r.raise_for_status()
        data=r.json()
        translated="".join(chunk[0] for chunk in data[0] if chunk[0])
        return translated if translated else text
    except: return text

def is_mostly_english(text:str)->bool:
    if not text: return False
    total=sum(c.isalpha() for c in text)
    if total==0: return False
    ascii_count=sum(c.isascii() and c.isalpha() for c in text)
    return (ascii_count/total)>0.85

# ==========================
# ROBUST: disable augmentation inside model
# ==========================
def disable_augmentation_layers(model):
    disabled=[]
    def identity_call(self,inputs,training=False): return inputs
    def walk(layer):
        lname=getattr(layer,"name","").lower()
        cname=layer._class.name_.lower()
        if "augmentation" in lname or any(k in cname for k in["random","flip","rotate","rotation","zoom","contrast","crop","augment","preprocess"]):
            try:
                layer.call=MethodType(identity_call,layer)
                layer.trainable=False; disabled.append(layer.name)
            except: pass
        if hasattr(layer,"layers"):
            for l in layer.layers: walk(l)
    walk(model); return model

# ==========================
# PAGE CONFIG
# ==========================
st.set_page_config(page_title="🌾FARMDOC",layout="wide",initial_sidebar_state="expanded")

# ==========================
# PREVENT SESSION-STATE CRASHES
# ==========================
if "report_text" not in st.session_state: st.session_state.report_text=""
if "auto_refresh_on" not in st.session_state: st.session_state.auto_refresh_on=True

# ==========================
# APP-WIDE STYLES
# ==========================
def set_background_and_styles():
    st.markdown("""
    <style>
    .stApp{background:linear-gradient(135deg,#0f1724 0%,#162033 45%,#20324a 100%) fixed;color:#e6eef8;}
    .block-container{background:rgba(10,14,20,0.45);border-radius:14px;padding:28px;box-shadow:0 8px 30px rgba(0,0,0,0.45);}
    h1,h2,h3,h4,h5,h6,p,div,span,label{color:#eef6ff!important;}
    .card{background:linear-gradient(180deg,rgba(255,255,255,0.02),rgba(255,255,255,0.01));border-radius:12px;padding:14px;margin-bottom:12px;box-shadow:0 6px 18px rgba(2,6,23,0.6);}
    .caption{font-size:12px;color:#d6e8ff;opacity:.8;}
    .stButton>button{background:linear-gradient(90deg,#2fb86f,#35c06f)!important;color:white!important;font-weight:600;border-radius:12px!important;}
    .stDownloadButton>button{background:rgba(255,255,255,0.06)!important;color:white!important;border-radius:10px!important;}
    video{transform:scaleX(-1)!important;}
    </style>""",unsafe_allow_html=True)
set_background_and_styles()

# ==========================
# PAGE HEADER
# ==========================
header_col1,header_col2=st.columns([0.9,0.1])
with header_col1:
    st.markdown("<h1 style='margin:0;'>🌱 FarmDoc</h1>",unsafe_allow_html=True)
    st.markdown("<div class='caption'>Detect plant disease, view live farm sensor data, and get an easy-to-follow treatment report in multiple languages.</div>",unsafe_allow_html=True)
st.markdown("---")

# ==========================
# LOAD MODEL
# ==========================
MODEL_URL="https://raw.githubusercontent.com/horrorinpeace/DISEASE-DETECTION/main/fix.h5"
MODEL_PATH="fix.h5"

@st.cache_resource
def load_model():
    if not os.path.exists(MODEL_PATH):
        os.makedirs("models",exist_ok=True)
        urllib.request.urlretrieve(MODEL_URL,MODEL_PATH)
    m=tf.keras.models.load_model(MODEL_PATH,compile=False)
    return disable_augmentation_layers(m)

try:
    model=load_model()
    CLASS_NAMES=[
        'CURRY POWDERY MILDEW','HEALTHY MILLET','HEALTHY POTATO','HEALTHY RICE','HEALTHY SUGARCANE','HEALTHY TEA LEAF',
        'HEALTHY TOMATO','HEALTHY WHEAT','MILLETS BLAST','MILLETS RUST','POTATO EARLY BLIGHT','POTATO LATE BLIGHT',
        'RICE BACTERIAL BLIGHT','RICE BROWN SPOT','RICE LEAF SMUT','SUGARCANE RED ROT','SUGARCANE RUST',
        'SUGARCANE YELLOW','TEA GRAY BLIGHT','TEA GREEN MIRID BUG','TEA HELOPELTIS','TOMATO LEAF MOLD',
        'TOMATO MOSAIC VIRUS','TOMATO SEPTORIA LEAF SPOT','WHEAT BROWN RUST','WHEAT LOOSE SMUT','WHEAT YELLOW RUST'
    ]
except Exception as e: st.error(f"Model failed to load: {e}"); model=None; CLASS_NAMES=[]

# ==========================
# READ KEY (NEW)
# ==========================
READ_KEY = "SO5QAU5RBCQ15WKD"

# ==========================
# FIXED — MERGE VALUES FROM MULTIPLE ENTRIES
# ==========================
def fetch_sensor_data():
    url=f"https://api.thingspeak.com/channels/3152731/feeds.json?api_key={READ_KEY}&results=40"
    try:
        res=requests.get(url,timeout=5).json()
        feeds=res.get("feeds",[])

        def latest(field):
            for row in reversed(feeds):
                if row.get(field) not in (None,""):
                    return row.get(field)
            return "None"

        timestamp = feeds[-1].get("created_at","—") if feeds else "—"

        return {
            "temperature":latest("field1"),
            "humidity":latest("field2"),
            "soil_moisture":latest("field3"),
            "air_quality":latest("field4"),
            "light_intensity":latest("field5"),
            "pressure":latest("field6"),
            "soil_temperature":latest("field7"),
            "timestamp":timestamp
        }
    except:
        return {"temperature":"None","humidity":"None","soil_moisture":"None",
                "air_quality":"None","light_intensity":"None","pressure":"None",
                "soil_temperature":"None","timestamp":"—"}

# ==========================
LANGUAGE_OPTIONS={
    "English":"English","हिन्दी (Hindi)":"Hindi","বাংলা (Bengali)":"Bengali","தமிழ் (Tamil)":"Tamil",
    "తెలుగు (Telugu)":"Telugu","ಕನ್ನಡ (Kannada)":"Kannada","മലയാളം (Malayalam)":"Malayalam",
    "मराठी (Marathi)":"Marathi","ગુજરાતી (Gujarati)":"Gujarati","ਪੰਜਾਬੀ (Punjabi)":"Punjabi",
    "ଓଡ଼ିଆ (Odia)":"Odia","اردو (Urdu)":"Urdu"
}

# ==========================
# SIDEBAR
# ==========================
st.sidebar.title("Menu")
page=st.sidebar.radio("Go to",["About","AI Detection Panel"])
st.sidebar.markdown("---")
st.sidebar.markdown("### Settings")
selected_language_display=st.sidebar.selectbox("Report language",list(LANGUAGE_OPTIONS.keys()))
st.session_state.selected_language=LANGUAGE_OPTIONS[selected_language_display]
api_key=st.sidebar.text_input("🔐 Groq API key",type="password")

# ==========================
# ABOUT PAGE
# ==========================
if page=="About":
    st.header("About FarmDoc AI")
    st.markdown("""
    <div class='card'>
    <strong>FarmDoc AI</strong> helps farmers detect plant diseases using a phone camera or uploaded images.
    It gives simple, clear advice on what the disease is, how it affects the crop, what actions to take,
    and how to prevent it — available in different languages.
    </div>""",unsafe_allow_html=True)

    st.markdown("### How it works")
    st.markdown("""
    1. Take a photo of the leaf or upload an image.
    2. The AI recommends the most likely disease.
    3. You generate the report.
    4. Download it and share.
    """)

# ==========================
# AI PANEL
# ==========================
elif page=="AI Detection Panel":

    st.header("Step 1 — Capture or Upload Plant Image")
    st.markdown("<div class='card'>Use your phone camera or upload a clear photo.</div>",unsafe_allow_html=True)

    uploaded_file=st.camera_input("📸 Take a photo")
    if uploaded_file is None:
        uploaded_file=st.file_uploader("Or upload",type=["png","jpg","jpeg"])

    if uploaded_file:
        image=Image.open(uploaded_file).convert("RGB")
        st.image(image,width=300)

        if model:
            img=image.resize((224,224))
            arr=tf.keras.preprocessing.image.img_to_array(img)
            arr=np.expand_dims(arr,0)
            pred=model.predict(arr)
            st.session_state.predicted_class=CLASS_NAMES[np.argmax(pred)]
            st.success(f"🌿 Detected: {st.session_state.predicted_class}")

    # ==========================
    # LIVE DATA
    # ==========================
    st.header("Step 2 — Live Farm Data")
    if st.session_state.auto_refresh_on:
        st_autorefresh(interval=5000,limit=None,key="sensor_refresh")

    data=fetch_sensor_data()
    c1,c2,c3,c4,c5,c6,c7 = st.columns(7)

    c1.metric("🌡 Temperature",f"{data['temperature']} °C")
    c2.metric("💧 Humidity",f"{data['humidity']} %")
    c3.metric("🌱 Soil Moisture",f"{data['soil_moisture']} %")
    c4.metric("🫁 Air Quality",f"{data['air_quality']} AQI")
    c5.metric("💡 Light Intensity",f"{data['light_intensity']} lx")
    c6.metric("🌬 Pressure",f"{data['pressure']} hPa")
    c7.metric("🌡 Soil Temp",f"{data['soil_temperature']} °C")

    st.caption(f"Last updated: {data['timestamp']}")

    # ==========================
    # REPORT GENERATION — UPDATED PROMPT
    # ==========================
    st.header("Step 3 — Get Farm Report")
    st.markdown("<div class='card'>The AI will write the report in selected language.</div>",unsafe_allow_html=True)

    if st.button("🧾 Generate Farm Report"):
        st.session_state.report_text=""
        if not api_key: st.error("Enter Groq API key.")
        elif not uploaded_file: st.error("Upload or capture image.")
        elif model is None: st.error("Model not loaded.")
        elif "predicted_class" not in st.session_state: st.error("No disease prediction.")
        else:
            st.session_state.auto_refresh_on=False
            try:
                with st.spinner("Writing report..."):

                    ###################################################
                    ##  🔥 NEW PROMPT INSERTED — ONLY ADDITIONS MADE  ##
                    ###################################################
                    prompt = f"""
                    You are a helpful agricultural assistant for farmers.

                    Write a VERY detailed, step-by-step farm advisory report in a simple way for farmers to understand in {st.session_state.selected_language}.

                    STRICT RULES (must follow all):
                    - Never skip information or stay vague.
                    - Always give specific names of fungicides/pesticides (generic name + 1–2 common brand examples if possible).
                    - ALWAYS give exact dose in:
                      • ml or g per litre of water
                      • ml or g per 15 L knapsack sprayer
                      • total quantity per acre (or per hectare) and approximate water volume.
                    - Clearly mention:
                      • how many times to spray
                      • gap between sprays (in days)
                      • waiting period before harvest, if needed.
                    - Clearly list ALL tools and materials needed:
                      • sprayer type + nozzle type
                      • measuring tools
                      • gloves, goggles, face mask, etc.
                    - Fill every bullet completely.

                    🔥 Live Farm Conditions to Consider While Writing Report:
                    Temperature = {data['temperature']} °C
                    Humidity = {data['humidity']} %
                    Soil Moisture = {data['soil_moisture']} %
                    Soil Temperature = {data['soil_temperature']} °C
                    Air Quality Index = {data['air_quality']}
                    Light Intensity = {data['light_intensity']} lx
                    Atmospheric Pressure = {data['pressure']} hPa

                    Based on these values, also include:
                    - Rain/Dry weather prediction using pressure
                    - Fungal disease risk if humidity is high
                    - Sunlight recommendations if light intensity is low
                    - Irrigation & mulching advice based on soil temperature
                    - Air quality affect on crops & farmer health
                    - Preventive action plan for next 7 days

                    Use THIS EXACT FORMAT:

                    - Disease Name:
                    - What It Means:
                    - Cause:
                    - Name of spray to be used & Amount to be sprayed:
                    - Tools and Materialy Step Process For Treatment (with exact measurements and timing):
                    - How many times to spray & gap between sprays:
                    - Safety Precautions for Farmers:
                    - Prevention Tips Needed (with quantities):
                    - Step By Step Process For Treatment (with exact measurements and timing):
                    - How many times to spray & gap between sprays:
                    - Safety Precautions for Farmers:
                    - Prevention Tips:
                    """

                    url="https://api.groq.com/openai/v1/chat/completions"
                    headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"}
                    payload={
                        "model":"meta-llama/llama-4-scout-17b-16e-instruct",
                        "messages":[{"role":"system","content":"You give farm advice."},{"role":"user","content":prompt}],
                        "temperature":0.6,"max_completion_tokens":800,"top_p":1,"stream":False
                    }

                    r=requests.post(url,headers=headers,json=payload,timeout=40)
                    output=r.json()["choices"][0]["message"]["content"]
                    st.session_state.report_text=(
                        translate_with_google(output,st.session_state.selected_language)
                        if st.session_state.selected_language!="English" and is_mostly_english(output)
                        else output
                    )
                    st.success("Report generated!")
            except Exception as e: st.error(f"Error: {e}")
            finally: st.session_state.auto_refresh_on=True

# ==========================
# SHOW REPORT + DOWNLOAD
# ==========================
if st.session_state.report_text:
    st.markdown("### 🌿 Your Farm Report")
    st.markdown(f"<div class='card'><pre style='white-space:pre-wrap'>{st.session_state.report_text}</pre></div>",unsafe_allow_html=True)
    st.download_button("📥 Download TXT",st.session_state.report_text.encode(),file_name="farm_report.txt")

    try:
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO

        d=Document(); d.add_heading("Farm Report",level=1)
        for line in st.session_state.report_text.splitlines():
            p=d.add_paragraph(line)
            for run in p.runs: run.font.size=Pt(12)

        if 'uploaded_file' in locals():
            d.add_page_break(); d.add_paragraph("Attached leaf image:")
            d.add_picture(BytesIO(uploaded_file.getbuffer()),width=Pt(300))

        buf=BytesIO(); d.save(buf); buf.seek(0)
        st.download_button("📥 Download DOCX",buf.read(),"farm_report.docx")
    except: st.warning("DOCX export not available.")

st.markdown("---")
st.markdown("<div class='caption'>FarmDoc © 2025 — Helping Farmers Grow Smarter</div>",unsafe_allow_html=True)
